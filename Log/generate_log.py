"""Generate EE297A Meeting Log File as .docx matching the reference format."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ── Cover / Header ──
h1 = doc.add_heading('EE297A - Graduate Meeting Log File', level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph('High-Performance Multimodal Transformer Inference on Commodity GPUs')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.size = Pt(14)

for line in [
    'Advisor: Prof. Chang Choo',
    '',
    'A Project Proposal Presented to',
    'The faculty of the Department of Electrical Engineering',
    'San José State University',
    '',
    'In Partial Fulfillment of the Requirements for the Degree',
    'Master of Science',
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph('By')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
for line in [
    'Jay Darji & 018180851',
    'EE297A Section 01, Spring 2026',
    'jaykumar.darji@sjsu.edu, +1 (669)-340-6428',
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
for line in [
    'Karma Patel & 018223517',
    'EE297A Section 01, Spring 2026',
    'karma.patel@sjsu.edu, +1 (510)-766-9741',
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
for line in [
    'Department of Electrical Engineering',
    'Charles W. Davidson College of Engineering',
    'San José State University',
    'San Jose, CA 95192-0084',
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── Meeting entries ──
meetings = [
    {
        "title": "Meeting 1 — January 24, 2026",
        "time": "08:52 AM – 10:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "This was the semester's first official introductory meeting. The focus was largely organizational, "
            "with Professor Choo confirming attendance, finalizing group arrangements, and going over expectations "
            "for graduate students. Jay and I confirmed our group and briefly introduced our initial idea: "
            "studying efficient transformer inference on memory-constrained hardware.\n\n"
            "Even though no technical work was expected yet, it was clear from the advisor's tone that graduate "
            "projects need consistent weekly movement. He emphasized maintaining logs, providing updates even when "
            "progress feels small, and being proactive about literature review and implementation planning."
        ),
        "actions": [
            "Start gathering background material on efficient transformer inference",
            "Discuss with Jay how we want to divide early tasks",
        ],
        "connection": "This meeting set the framework and aligned our expectations with the structure we eventually outlined in our proposal.",
    },
    {
        "title": "Meeting 2 — January 31, 2026",
        "time": "08:52 AM – 10:30 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "This week was our second introductory meeting. Choo checked subgroup placement and we confirmed "
            "our meeting cycle. The rest of the meeting largely consisted of listening to other teams' project ideas "
            "and receiving feedback on the scope of work expected for a master's project.\n\n"
            "Jay and I shared notes privately during the meeting. We discussed our target hardware (NVIDIA RTX 3050 Ti "
            "with 4GB VRAM) and realized we needed to carefully select model candidates that would fit this strict budget. "
            "We also discussed setting up a reliable GPU environment with PyTorch and CUDA. Choo emphasized that we must "
            "document issues encountered, not just final results."
        ),
        "actions": [
            "Finalize model shortlist within 4GB VRAM constraint (Qwen2.5-3B, Phi-2, etc.)",
            "Begin PyTorch and CUDA environment setup",
            "Draft our initial VRAM budget worksheet",
        ],
        "connection": "This meeting helped us understand the depth required and directly led to our Week 01 and 02 planning phases.",
    },
    {
        "title": "Meeting 3 — February 7, 2026",
        "time": "09:00 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We presented our first technical progress. We successfully configured CUDA 12.4 and PyTorch, and "
            "demonstrated our automated setup validation script (week02_setup_validation.py). We also ran our first "
            "baseline tests using Qwen2.5-3B-Instruct with 4-bit NF4 quantization.\n\n"
            "We built a benchmarking harness to measure p50/p95 latency, throughput, and peak VRAM. We noticed that "
            "VRAM usage spiked significantly during attention layers at longer sequence lengths. Our initial tests "
            "showed an OOM (Out Of Memory) threshold around 6,784 tokens with standard eager attention.\n\n"
            "Choo liked the systematic benchmarking approach and suggested we look deeply into PyTorch profiler traces "
            "to confirm exactly where the time and memory are being spent."
        ),
        "actions": [
            "Implement PyTorch profiler wrapper to capture kernel traces",
            "Run bottleneck analysis on the eager baseline",
        ],
        "connection": "This meeting validated our environment setup and established the baseline for all future optimizations.",
    },
    {
        "title": "Meeting 4 — February 14, 2026",
        "time": "08:58 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We presented the results of our Week 04 profiling. Using the PyTorch profiler, we analyzed the CUDA "
            "traces and confirmed that attention-related kernels (like scaled dot-product) dominated the execution time, "
            "especially as sequence length grew. We also documented the availability of different attention backends "
            "on our specific hardware.\n\n"
            "We explained our plan to evaluate PyTorch's native SDPA (Scaled Dot-Product Attention) and to explicitly "
            "install and test the FlashAttention-2 package from source or pre-built wheels, as Flash SDP wasn't natively "
            "available in our PyTorch Windows wheel.\n\n"
            "Choo agreed with this direction and reminded us to ensure we verify correctness (top-1 token agreement) "
            "when switching attention backends."
        ),
        "actions": [
            "Implement Phase 1 and 2 attention tests (Eager and SDPA)",
            "Install FlashAttention-2 explicitly and integrate it as Phase 3",
            "Set up correctness logging using fixed prompt greedy decoding",
        ],
        "connection": "This meeting shifted our focus from measurement to active optimization, targeting the identified attention bottlenecks.",
    },
    {
        "title": "Meeting 5 — February 21, 2026",
        "time": "09:00 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We shared our initial attention optimization integration (Week 05). We successfully implemented "
            "scripts to test Eager, SDPA default, SDPA math, and explicit FlashAttention-2. "
            "We encountered some stability issues; for example, memory-efficient SDPA failed on GQA models due to "
            "head count mismatches, but standard SDPA and FA2 worked beautifully.\n\n"
            "We ran correctness checks across all working backends and achieved 100% top-1 token agreement with the "
            "baseline, proving that our optimizations didn't degrade the model's deterministic output.\n\n"
            "Choo was pleased with the rigorous correctness checking and asked us to prepare final tradeoff tables "
            "showing how much extra context length we gained before hitting OOM."
        ),
        "actions": [
            "Run complete benchmark sweeps for all working attention configs",
            "Measure new OOM thresholds for SDPA and FlashAttention-2",
        ],
        "connection": "This meeting proved that attention optimization was feasible and correct on our commodity hardware.",
    },
    {
        "title": "Meeting 6 — February 28, 2026",
        "time": "08:56 AM – 10:45 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We presented the finalized attention results (Week 06). The improvements were significant: "
            "FlashAttention-2 pushed our OOM threshold from 6,784 tokens up to 8,064 tokens (+18.9%) while "
            "maintaining perfect correctness. SDPA default provided a smaller, zero-effort boost to 7,040 tokens.\n\n"
            "We shared our generated CSV/JSON tables detailing p50 latency, throughput, and VRAM at prompt lengths "
            "of 128, 512, and 1024.\n\n"
            "With attention optimized, Choo directed us to begin looking at the KV-cache, as it is the next major "
            "memory bottleneck during long-context autoregressive generation."
        ),
        "actions": [
            "Research KV-cache quantization methods (e.g., optimum-quanto)",
            "Build a prototype to run inference with INT4 and INT2 KV-cache",
        ],
        "connection": "This meeting concluded the attention phase of our project and initiated the KV-cache optimization phase.",
    },
    {
        "title": "Meeting 7 — March 7, 2026",
        "time": "09:00 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We demonstrated our Week 07 KV-cache quantization prototype. We successfully integrated "
            "optimum-quanto via the transformers QuantizedCacheConfig API, enabling INT4 and INT2 KV-cache "
            "on the fly. \n\n"
            "We measured the memory footprint per token and found noticeable VRAM savings. For INT4, we "
            "reduced per-token memory overhead by approximately 13%. However, we noted a slight increase "
            "in latency due to the quantization/dequantization overhead on each forward pass.\n\n"
            "Choo advised that memory savings are only useful if the model retains its quality, and told us "
            "to implement zero-shot quality evaluation on standard academic datasets to verify INT4's viability."
        ),
        "actions": [
            "Implement quality evaluation harness for ARC-Easy, BoolQ, and HellaSwag",
            "Run full tradeoff analysis: memory saved vs. latency vs. quality retention",
        ],
        "connection": "This meeting proved the feasibility of KV-cache quantization and set up the quality validation step.",
    },
    {
        "title": "Meeting 8 — March 14, 2026",
        "time": "08:58 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We presented our KV-cache experiments and quality evaluation (Week 08). We ran 0-shot log-likelihood "
            "evaluations on 200 examples each from ARC-Easy, BoolQ, and HellaSwag.\n\n"
            "The results showed that INT4 KV-cache maintained accuracy within ≤2% of the FP16 baseline across all "
            "three datasets. INT2 showed more significant degradation. We decided to select INT4 as our final KV "
            "configuration, as it offered meaningful VRAM savings without unacceptable latency or quality penalties.\n\n"
            "Choo agreed with our selection of INT4 and challenged us to see if we could combine FlashAttention-2 "
            "with INT4 KV-cache in the same run."
        ),
        "actions": [
            "Test combined configurations: SDPA + INT4 and FA2 + INT4",
            "Document any framework incompatibilities",
        ],
        "connection": "This meeting finalized our KV-cache strategy and initiated the complex combined configuration testing.",
    },
    {
        "title": "Meeting 9 — March 21, 2026",
        "time": "09:00 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We reported on our Week 09 combined configuration testing. We found that combining FA2 and INT4 KV-cache "
            "worked perfectly on MHA (Multi-Head Attention) models like Phi-2 and Llama-2-7B. However, on GQA "
            "models like Qwen2.5-3B, we hit framework-level assertion errors related to head-count mismatches when "
            "using FlashAttention alongside quantized caches.\n\n"
            "We documented this incompatibility thoroughly in our logs and generated a final benchmark runbook that "
            "intelligently skips unstable combinations while benchmarking the working ones.\n\n"
            "Choo appreciated the deep dive into framework limitations and told us to execute the final full benchmark "
            "suite for all working configurations."
        ),
        "actions": [
            "Run the full Week 10 benchmark suite using the finalized runbook",
            "Consolidate all results into a single comprehensive table",
        ],
        "connection": "This meeting finalized the exact configurations that would make it into our final report.",
    },
    {
        "title": "Meeting 10 — March 28, 2026",
        "time": "08:56 AM – 10:45 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We presented our Week 10 Full Benchmark Suite. The automated pipeline successfully benchmarked the "
            "baseline, SDPA, FA2, and INT4 KV configurations across multiple prompt lengths. We generated a large "
            "consolidated CSV/JSON dataset containing thousands of data points on latency, throughput, and VRAM.\n\n"
            "From this data, we drafted a 'Configuration Guidelines' document for 4GB VRAM laptops, explicitly "
            "recommending when to prioritize FA2 (for max context) versus INT4 KV (for VRAM savings at generation).\n\n"
            "Choo was very impressed with the consolidated data. He asked if these techniques would also apply "
            "to multimodal vision-language models, which have much larger input sequences due to image tokens."
        ),
        "actions": [
            "Integrate multimodal models (Phi-3.5-vision and LLaVA-1.5-7B) into the codebase",
            "Download test images and prepare image+text prompt benchmarking",
        ],
        "connection": "This meeting marked the completion of the core text-only project and initiated the multimodal extension.",
    },
    {
        "title": "Meeting 11 — April 4, 2026",
        "time": "09:00 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We introduced our new `multimodal_loader.py` module. We demonstrated loading Phi-3.5-vision and "
            "LLaVA-1.5-7B within our 4GB VRAM budget. LLaVA-1.5-7B at 4-bit quantization sat right at ~3.5GB, "
            "leaving very little headroom, which made optimizations critical.\n\n"
            "We showed that we could successfully apply SDPA, explicit FA2, and KV-cache quantization to these "
            "multimodal models. We ran initial inference tests with a sample image, showing the model correctly "
            "describing objects and colors.\n\n"
            "Choo was highly satisfied that our architecture extended gracefully to multimodal workloads."
        ),
        "actions": [
            "Run the full benchmark suite on the multimodal models",
            "Capture sample generation texts for qualitative review",
        ],
        "connection": "This meeting proved the extensibility of our optimization pipeline to complex vision-language tasks.",
    },
    {
        "title": "Meeting 12 — April 11, 2026",
        "time": "08:58 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We presented the complete multimodal benchmark results. The trends observed in text-only models "
            "held true: FA2 significantly reduced VRAM spikes during the massive prompt processing phase (which "
            "includes hundreds of image tokens), preventing OOM on LLaVA-1.5-7B.\n\n"
            "We also reviewed the generated sample texts to ensure the multimodal models weren't hallucinating "
            "under quantization. The qualitative output remained robust.\n\n"
            "With all experiments complete, Choo instructed us to focus entirely on documentation, plotting, and "
            "report writing."
        ),
        "actions": [
            "Write `generate_plots.py` to create publication-quality figures",
            "Begin drafting the final LaTeX/HTML project report",
        ],
        "connection": "This meeting concluded all experimental data gathering for the semester.",
    },
    {
        "title": "Meeting 13 — April 18, 2026",
        "time": "09:00 AM – 11:00 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "We focused heavily on presentation. We showed Choo the plots generated by our new script, "
            "including 'Latency vs. Sequence Length', 'Peak VRAM scaling', and 'OOM Threshold Comparisons'. "
            "The visual data clearly communicated the 18.9% context length gain from FA2 and the VRAM reduction "
            "from INT4 KV-cache.\n\n"
            "We reviewed an early draft of our project report. Choo provided feedback on formatting, suggesting "
            "we emphasize our 'commodity hardware' angle more strongly in the abstract, as achieving these results "
            "on a 4GB RTX 3050 Ti laptop is a strong engineering achievement."
        ),
        "actions": [
            "Finalize the `project_report.html` and `report.tex` documents",
            "Create a `PROJECT_EXPLANATION_AND_QA.md` file to prepare for technical defense",
        ],
        "connection": "This meeting was critical for refining how we communicate our 13 weeks of technical work.",
    },
    {
        "title": "Meeting 14 — April 25, 2026",
        "time": "08:56 AM – 10:45 AM",
        "attendees": "Karma Patel, Jaykumar Darji, Professor Choo",
        "body": (
            "This was our final meeting of the semester. We presented the fully completed project repository, "
            "including all source code, benchmark data, plots, and final reports. We demonstrated how a user "
            "could reproduce our entire semester's work using a single `python scripts/run_all.py` command.\n\n"
            "We discussed our QA preparation document, going over potential defense questions regarding "
            "NF4 quantization mathematics and FlashAttention memory I/O.\n\n"
            "Choo commended the rigor, reproducibility, and organization of our project. He approved our final "
            "deliverables and confirmed we were ready for our final submission and presentation."
        ),
        "actions": [
            "Push all final code, docs, and the Q&A file to the main GitHub repository",
            "Submit the final report to the department",
            "Generate this final Meeting Log File",
        ],
        "connection": "This meeting successfully closed out the project, validating all deliverables.",
    },
]

for m in meetings:
    h = doc.add_heading(m["title"], level=2)

    doc.add_paragraph()
    doc.add_paragraph(f"Time: {m['time']}")
    doc.add_paragraph(f"Attendees: {m['attendees']}")
    doc.add_paragraph()

    for para_text in m["body"].split('\n\n'):
        doc.add_paragraph(para_text.strip())

    doc.add_paragraph("Action Items:", style='Normal').runs[0].bold = True
    for item in m["actions"]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    cp = doc.add_paragraph("Connection to Proposal:", style='Normal')
    cp.runs[0].bold = True
    doc.add_paragraph(m["connection"])
    doc.add_paragraph()

# Save
out_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "018223517 018180851 EE297A Spring 2026 Meeting Log File_v2.docx",
)
doc.save(out_path)
print(f"Saved to: {out_path}")
